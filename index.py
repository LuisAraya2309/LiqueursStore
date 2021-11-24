#Libraries
from flask import Flask,render_template,request
from flask.templating import render_template_string
from requests.models import parse_url #Flask framework
from databaseConnection import *
from auxiliarFunctions import *
import json ,requests, os,random
import shutil
from geopy.geocoders import Nominatim 
import datetime

from pymongo import MongoClient
from docx import Document
import os
#App creation
app = Flask(__name__)


#Global variables
username =""
api = "http://localhost:3001/"
imageSrc = ""
country="Argentina"
'''
newRequest = api+'users'
resp = requests.get(newRequest)
p = resp.json()
'''

#Routes
@app.route('/',methods=['GET'])   #Login page
def login():
    return render_template('login.html')

@app.route('/signIn',methods=['GET','POST'])  #Sign in
def validateUser():
    
    countrySelected = request.form['countrySelected']
    global country
    country = countrySelected
    user = request.form['user']
    password = request.form['password']
    
    #Get subsidiaries
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'SELECT S.Title FROM dbo.Subsidiary AS S'
            cursor.execute(querySubisidiary)
            queryResult = cursor.fetchall()
            subsidiaries = []
            for subsidiary in queryResult:
                subsidiaries.append(subsidiary[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    #Get Products
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            queryLiqueurs = 'SELECT L.Title FROM dbo.Liqueurs AS L'
            cursor.execute(queryLiqueurs)
            queryResult = cursor.fetchall()
            products = []
            for product in queryResult:
                products.append(product[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()  
    
    
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            query = 'EXEC sp_SignIn ? , ? , ?'
            cursor.execute(query,(user,password,0))
            queryResult = cursor.fetchall()
            interfaces = {'Consultant':'consultant.html','Administrator':'admin.html'}
            validUser = queryResult[0][0]
            userType = queryResult[0][1]

            if validUser != 1:
                global username
                username = user
                return render_template(interfaces[userType],**locals())
            else:
                return render_template('login.html') + '''<div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">Usuario o contraseña inválido. Vuelva a intentarlo.
                                    <a href="/beginSignUp">Registrarse</a></div>
                                    <div class="content-buttons"><a href="#" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/";
                                        });
                            </script>
                            '''

    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()
 
@app.route('/subSchedules',methods=['GET','POST'])
def showSchedules():
    subsidiary = request.form['subsidiary']
    dbConnection = connectToDatabase(country)
    optionSelect = ""
    
    
    try:
        with dbConnection.cursor() as cursor:
            query = 'EXEC sp_ViewSchedule ? , ?'
            cursor.execute(query,(subsidiary,0))
            queryResult = cursor.fetchall()
            schedulesDay = []
            for day in queryResult:
                optionSelect = ""
                hour = str(day[2]).split(':')
                hour = hour[0] +":"+ hour[1]
                optionSelect+=(hour + "am  -  ")
                hour = str(day[3]).split(':')
                hour = hour[0] +":"+ hour[1]
                optionSelect+=(hour+"pm.")
                schedulesDay.append(optionSelect)
            return render_template('schedules.html',**locals())
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()

        
@app.route('/queryProduct',methods=['GET','POST'])
def queryProduct():
    productName = request.form['productName']
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            query = 'EXEC sp_ViewProduct ? , ?'
            cursor.execute(query,(productName,0))
            queryResult = cursor.fetchall()
            simpleResult = queryResult[0]
            combinations = "Combinaciones sugeridas: "
            for result in queryResult:
                combinations+=(str(result[6]) + ', ')
            combinations = combinations[:-2]+'.'
            price = int(simpleResult[1])
            agedType = str(simpleResult[2])
            years = str(simpleResult[3])
            origin = str(simpleResult[4])
            tempFile = open('temp.jpg','wb')
            tempFile.write(queryResult[0][5])
            tempFile.close()
            shutil.move('temp.jpg','./static/images/temp.jpg')
            nearestSubsidiares = subsidiariesNear(username,productName,country)
            nearestSub = nearestSubsidiares[0][1]
            distance = int(nearestSubsidiares[0][2])
            distance = str(distance/1000)[:-1]
            nearestSub+=(" a solamente " + str(distance))
            subsAvailableList = subsAvailable(productName,country)
            availablesSubs = "Disponible en: "
            
            if len(subsAvailableList) != 1:
                for result in subsAvailableList:
                    availablesSubs+=(result + ', ')
                availablesSubs = availablesSubs[:-2]+'.'
            else:
                availablesSubs += subsAvailableList[0] + '.'


            #destiny =   str(os.path.realpath('temp.jpg')) 
            #return str(destiny)
            #os.startfile(destiny)

            return render_template('productQuery.html',**locals())



    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        

@app.route('/changeStock',methods=['GET','POST'])
def changeStock():
    subsidiary = request.form['subsidiary']
    product = request.form['product']
    stock = request.form['stock']
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            query = 'EXEC sp_UpdateStock ? , ?, ?, ?'
            cursor.execute(query,(product,stock,subsidiary,0))
            queryResult = cursor.fetchall()
            resultCode = int(queryResult[0][0])
            if resultCode == 0:
                return  '''
                        <!DOCTYPE html>
                        <html>
                            <head>
                                <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                                <link href="./static/css/admin.css" rel="stylesheet" />
                                <title>Ventana emergente</title>
                            </head>
                            <div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">El stock se ha aumentando correctamente</div>
                                    <div class="content-buttons"><a href="#" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/returnAdmin";
                                        });
                            </script>
                            '''
                            
            else:
                '''
                        <!DOCTYPE html>
                        <html>
                            <head>
                                <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                                <link href="./static/css/admin.css" rel="stylesheet" />
                                <title>Modificar Inventario</title>
                            </head>
                            <div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">'''+"El producto a cambiar no existe."+'''
                                    </div>
                                    <div class="content-buttons"><a href="#" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/returnAdmin";
                                        });
                            </script>
                            '''
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()  

     
        
@app.route('/listEmployees',methods=['GET','POST'])
def listEmployees():
    subsidiary = request.form['subsidiaryEmployees']
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'EXEC sp_EmployeesPerSubsidiary ?,?'
            cursor.execute(querySubisidiary,(subsidiary,0))
            queryResult = cursor.fetchall()
            employees = []
            for result in queryResult:
                employees.append(str(result[0]))
            
            return render_template('listEmployees.html',**locals())    
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
        
@app.route('/showEmployee',methods=['GET','POST'])
def showEmployee():
    employeeSelected = request.form['employeeSelected']
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'EXEC sp_ViewEmployee ?,?'
            cursor.execute(querySubisidiary,(employeeSelected,0))
            queryResult = cursor.fetchall()
            result = queryResult[0]
            employeeName = result[2]
            email = result[3]
            phone = result[4]
            age = result[5]
            salary = str(result[-1]).split('.')[0] 
            job = result[-2]
            tempFile = open('tempE.jpg','wb')
            tempFile.write(result[6])
            tempFile.close()
            shutil.move('tempE.jpg','./static/images/tempE.jpg')
            
            
            
            return render_template('viewEmployee.html',**locals())    
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
    

@app.route('/startShopping',methods=['GET','POST'])
def startShopping():
    productName = request.form['productName']
    userName = username
    Country = country
    #Get subsidiaries
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'SELECT S.Title FROM dbo.Subsidiary AS S'
            cursor.execute(querySubisidiary)
            queryResult = cursor.fetchall()
            subsidiaries = []
            for subsidiary in queryResult:
                subsidiaries.append(subsidiary[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
    
    #Get subsidiaries
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'SELECT E.Title FROM dbo.Employees AS E WHERE E.IdEmployeeType = 1'
            cursor.execute(querySubisidiary)
            queryResult = cursor.fetchall()
            billers = []
            for biller in queryResult:
                billers.append(biller[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
        
    #Get price
    dbConnection = connectToDatabase(country)
    price =""
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'SELECT L.Price FROM dbo.Liqueurs AS L WHERE L.Title = ?'
            cursor.execute(querySubisidiary,(productName))
            queryResult = cursor.fetchall()
            price = str(int(queryResult[0][0]))
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    #Get discount
    dbConnection = connectToDatabase(country)
    discount =""
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'EXEC sp_CustomerDiscount ?, ?'
            cursor.execute(querySubisidiary,(username,0))
            queryResult = cursor.fetchall()
            hasDiscount =queryResult[0][0]
            finalDiscount =  (int(price)*0.10)
            discount = str(finalDiscount)  if hasDiscount == 'True' else "0.0"
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    return render_template('invoice.html',**locals())


@app.route('/buyProduct',methods=['GET','POST'])
def buyProduct():
    subsidiary = request.form['subsidiary']
    userName = request.form['userName']
    price = request.form['price'][1:]
    discount = request.form['discount'][1:]
    biller = request.form['biller']
    address=request.form['address']
    Country = request.form['Country']
    productName = request.form['productName']
    paymentType = request.form['paymentType']
    
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'EXEC sp_buyProduct ?, ?, ?, ?, ?, ?, ?, ?, ?'
            cursor.execute(querySubisidiary,(subsidiary,userName,price, discount, biller,address,productName,paymentType,0))
                
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'EXEC sp_CustomerEmail ?, ?'
            cursor.execute(querySubisidiary,(userName,0))
            email = cursor.fetchall()[0][0]
            message = "Su compra del artículo " + productName + " se ha realizado con éxito.\n"
            message += "Su envío se acaba de realizar, por favor esté atento en estos días a la llegada del producto."
            sendEmail(email,message,'Confirmación de compra')
            
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()  
    
    return  '''
            <!DOCTYPE html>
            <html>
                <head>
                    <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                    <link href="./static/css/admin.css" rel="stylesheet" />
                    <title>Mensaje emergente</title>
                </head>
                <div class="window-notice" id="window-notice" >
                    <div class="content">
                        <div class="content-text">Compra realizada con exito</div>
                        <div class="content-buttons"><a href="#" id="close-button">Aceptar</a></div>
                    </div>
                </div>
                <script>
                            let close_button = document.getElementById('close-button');
                                close_button.addEventListener("click", function(e) {
                                e.preventDefault();
                                document.getElementById("window-notice").style.display = "none";
                                window.location.href="#";
                            });
                </script>
                '''    
    
    
    


@app.route('/returnAdmin',methods=['GET','POST'])
def returnAdmin():
    #Get Products
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            queryLiqueurs = 'SELECT L.Title FROM dbo.Liqueurs AS L'
            cursor.execute(queryLiqueurs)
            queryResult = cursor.fetchall()
            products = []
            for product in queryResult:
                products.append(product[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    #Get subsidiaries
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'SELECT S.Title FROM dbo.Subsidiary AS S'
            cursor.execute(querySubisidiary)
            queryResult = cursor.fetchall()
            subsidiaries = []
            for subsidiary in queryResult:
                subsidiaries.append(subsidiary[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    return render_template('admin.html', **locals())
        
@app.route('/productStock',methods=['GET','POST'])
def productStock():
    productName = request.form['productName']
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            query = 'EXEC sp_ProductStock ?, ?'
            cursor.execute(query,(productName,0))
            queryResult = cursor.fetchall()
            subsidiaryStock=""
            
            for result in queryResult:
                subsidiaryStock+=result[1]
                subsidiaryStock+=": "
                subsidiaryStock+=(str(result[0])+"\n")
                
            return  '''
                        <!DOCTYPE html>
                        <html>
                            <head>
                                <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                                <link href="./static/css/admin.css" rel="stylesheet" />
                                <title>Modificar Inventario</title>
                            </head>
                            <div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">'''+subsidiaryStock+'''
                                    </div>
                                    <div class="content-buttons"><a href="/returnAdmin" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/returnAdmin";
                                        });
                            </script>
                            '''
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()  

            
@app.route('/editPrice',methods=['GET','POST'])
def editPrice():
    productName = request.form['productName']
    newPrice = request.form['newPrice']
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            query = 'EXEC sp_UpdatePrice ?, ?, ?'
            cursor.execute(query,(newPrice,productName,0))
            queryResult = cursor.fetchall()
            resultCode = int(queryResult[0][0])
            if resultCode == 0:
                return  '''
                        <!DOCTYPE html>
                        <html>
                            <head>
                                <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                                <link href="./static/css/admin.css" rel="stylesheet" />
                                <title>Modificar Inventario</title>
                            </head>
                            <div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">'''+"Precio editado con éxito."+'''
                                    </div>
                                    <div class="content-buttons"><a href="#" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/returnAdmin";
                                        });
                            </script>
                            '''
                            
            else:
                '''
                        <!DOCTYPE html>
                        <html>
                            <head>
                                <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                                <link href="./static/css/admin.css" rel="stylesheet" />
                                <title>Modificar Inventario</title>
                            </head>
                            <div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">'''+"El producto a cambiar no existe."+'''
                                    </div>
                                    <div class="content-buttons"><a href="#" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/returnAdmin";
                                        });
                            </script>
                            '''
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()  
    
@app.route('/returnAddProduct',methods=['GET','POST'])
def returnAddNewProduct():
    
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'SELECT S.Title FROM dbo.Subsidiary AS S'
            cursor.execute(querySubisidiary)
            queryResult = cursor.fetchall()
            subsidiaries = []
            for subsidiary in queryResult:
                subsidiaries.append(subsidiary[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    return render_template('addProduct.html',**locals())
    
  
@app.route('/addProduct',methods=['GET','POST'])
def addNewProduct():
    
    #Insert data
    title = request.form['Nombre']
    subsidiary = request.form['Subsidiaria']
    price = request.form['Precio']
    years = request.form['Años']
    origin = request.form['Origen']
    agedType = request.form['Añejado']
    amount = request.form['Cantidad']
    photo = request.files['Foto']
    food = request.form['Comida']
    
    if photo.filename != '':
        photo.save(photo.filename)
    fileName = photo.filename
    
    imageLiqueur = open(fileName,'rb')
    imageBytes = imageLiqueur.read()
    imageLiqueur.close()
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            query = 'EXEC sp_AddNewProduct ? , ? , ? ,? , ? , ?, ? , ? , ?, ?, ?'
            cursor.execute(query,(agedType,title,origin,price,years,amount,subsidiary,username,imageBytes,food,0))
            queryResult = cursor.fetchall()
            outResultCode = queryResult[0][0]
            imageLiqueur.close()
            if outResultCode != 1:
                return render_template('addProduct.html')+ '''<div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">El producto fue ingresado con exito.</div>
                                    <div class="content-buttons"><a href="/returnAdmin" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/returnAdmin";
                                        });
                            </script>
                            '''
            else:
                return render_template('addProduct.html') + '''<div class="window-notice" id="window-notice" >
                                <div class="content">
                                    <div class="content-text">Error al insertar el producto. Vuelva a intentarlo.</div>
                                    <div class="content-buttons"><a href="/returnAdmin" id="close-button">Aceptar</a></div>
                                </div>
                            </div>
                            <script>
                                        let close_button = document.getElementById('close-button');
                                            close_button.addEventListener("click", function(e) {
                                            e.preventDefault();
                                            document.getElementById("window-notice").style.display = "none";
                                            window.location.href="/returnAdmin";
                                        });
                            </script>
                            '''

    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()


def saleXCountryDoc(listTransaction):
    document = Document()
    document.add_heading(country + ' registró ' + str(len(listTransaction)) + ' venta(s)', 0)
    for transaction in listTransaction:
        document.add_paragraph("")
        p = document.add_paragraph("Sucursal: "+str(transaction[0]) + ", Fecha: "+ str(transaction[1]) + " , Cliente: " + str(transaction[2]) + " , Producto: " + str(transaction[3]), style = 'List Bullet')
        p.add_run(", Tipo de pago: "+ str(transaction[4])+ " , Descuento: " + str(transaction[5])+" , Precio final: " + str(transaction[6]))
        
    document.save('Reporte de ventas.docx')

def saleXSubsidiaryDoc(listTransaction, subsidiary):

    document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
    document.add_heading("La sucursal " + subsidiary + ' registró ' + str(len(listTransaction)) + ' venta(s)', level=1)
    for transaction in listTransaction:
        document.add_paragraph("")
        p = document.add_paragraph("Fecha: "+ str(transaction[0]) + " , Cliente: " + str(transaction[1]) + " , Producto: " + str(transaction[2]), style = 'List Bullet')
        p.add_run(", Tipo de pago: "+ str(transaction[3])+ " , Descuento: " + str(transaction[4])+" , Precio final: " + str(transaction[5]))
        
    document.save('Reporte de ventas.docx')
       
def saleXSubsidiaryXProductsDoc(listTransaction, productTitle):
    document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
    document.add_heading(str(len(listTransaction)) + ' venta(s) del licor '+ productTitle, level=1)
    for transaction in listTransaction:
        document.add_paragraph("")
        p = document.add_paragraph("Fecha: "+ str(transaction[0]) + " , Cliente: " + str(transaction[1]),  style = 'List Bullet')
        p.add_run(", Tipo de pago: "+ str(transaction[2])+ " , Descuento: " + str(transaction[3])+" , Precio final: " + str(transaction[4]))
        
    document.save('Reporte de ventas.docx')      

def saleXDateDoc(date, subsidiaryTitle, listTransaction):
    document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
    document.add_heading(subsidiaryTitle + ' registró '+str(len(listTransaction))+ ' venta(s) el día '+ str(date) , level=1)
    for transaction in listTransaction:
        document.add_paragraph("")
        p = document.add_paragraph("Cliente: " + str(transaction[2]) + " , Producto: " + str(transaction[3]),  style = 'List Bullet')
        p.add_run(", Tipo de pago: "+ str(transaction[4])+ " , Descuento: " + str(transaction[5])+" , Precio final: " + str(transaction[6]))
        
    document.save('Reporte de ventas.docx')    

def saleXPaymentTypeDoc(listTransaction):
    document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
    document.add_heading(str(listTransaction[0][0]) + ' registró '+str(len(listTransaction))+ ' venta(s) por medio de '+ str(listTransaction[0][4]), level=1)
    for transaction in listTransaction:
        document.add_paragraph("")
        p = document.add_paragraph("Sucursal: "+str(transaction[0]) + ", Fecha: "+ str(transaction[1]) + " , Cliente: " + str(transaction[2]) + " , Producto: " + str(transaction[3]), style = 'List Bullet')
        p.add_run(", Tipo de pago: "+ str(transaction[4])+ " , Descuento: " + str(transaction[5])+" , Precio final: " + str(transaction[6]))
        
    document.save('Reporte de ventas.docx')   

def saleXDatexPaymentTypeDoc(listTransaction):
    document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
    document.add_heading(str(listTransaction[0][0]) + ' registró '+str(len(listTransaction))+ ' venta(s) el día '+ str(listTransaction[0][1]) +' por medio de '+ str(listTransaction[0][4]), level=1)
    for transaction in listTransaction:
        document.add_paragraph("")
        p = document.add_paragraph("Sucursal: "+str(transaction[0]) + ", Fecha: "+ str(transaction[1]) + " , Cliente: " + str(transaction[2]) + " , Producto: " + str(transaction[3]), style = 'List Bullet')
        p.add_run(", Tipo de pago: "+ str(transaction[4])+ " , Descuento: " + str(transaction[5])+" , Precio final: " + str(transaction[6]))
        
    document.save('Reporte de ventas.docx')   

def createSalesXCountry():
    
    dbConnectionSalesXSubs = connectToDatabase(country)
    try:
        with dbConnectionSalesXSubs.cursor() as cursor:
            querySalesXCountry = 'EXEC sp_salesXCountry ?'
            cursor.execute(querySalesXCountry,(0))
            queryResult = cursor.fetchall()
            saleXCountryDoc(queryResult)
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnectionSalesXSubs.close()  

def subsidiaryXCountry(): 
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'EXEC sp_subsidiaryXCountry ?'
            cursor.execute(querySubisidiary,(0))
            queryResult = cursor.fetchall()
            subsidiaryList = []
            for i in range(0,len(queryResult)):
                subsidiaryList.append(queryResult[i][0])
            return subsidiaryList
                
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()       
        
def createSalesXSubisidiary(idsubsidiary, titleSubsidiary):
    
    dbConnectionSalesXSubs = connectToDatabase(country)
    try:
        with dbConnectionSalesXSubs.cursor() as cursor:
            querySalesXSubisidiary = 'EXEC sp_salesXSubsidiary ? ,?'
            cursor.execute(querySalesXSubisidiary,(idsubsidiary,0))
            queryResult = cursor.fetchall()
            saleXSubsidiaryDoc(queryResult, titleSubsidiary)
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnectionSalesXSubs.close()               

def productsXSubsidiary(idsubsidiary):
    dbConnection1 = connectToDatabase(country)
    try:
        with dbConnection1.cursor() as cursor:
            queryProductXSubisidiary = 'EXEC sp_productxSubsidiary ? ,?'
            cursor.execute(queryProductXSubisidiary,(idsubsidiary,0))
            queryResult = cursor.fetchall()
            productXSubsidiaryList = []
            for i in range(0,len(queryResult)):
                productXSubsidiaryList.append(queryResult[i][0]) 
            return productXSubsidiaryList 
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
        
    finally:
        dbConnection1.close()

def createSaleXSubsidiaryXProduct(idProduct, idSubsidiary, productTitle):
    
    dbConnection2 = connectToDatabase(country)
    try:
        with dbConnection2.cursor() as cursor:
            salesXSubsidiaryXProduct = 'EXEC sp_salesXSubsidiaryXProduct ? , ? , ?'
            cursor.execute(salesXSubsidiaryXProduct,(idProduct,idSubsidiary,0))
            queryResult = cursor.fetchall()
            salesXSubsidiaryXProductList = queryResult                              
            if salesXSubsidiaryXProductList :
                saleXSubsidiaryXProductsDoc(salesXSubsidiaryXProductList, productTitle)
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection2.close()   

def getDate(subsidiary):
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            queryGetDates = 'EXEC sp_getDates ? , ?'
            cursor.execute(queryGetDates,(subsidiary,0))
            queryResult = cursor.fetchall()
            dateList = []
            for i in range(0,len(queryResult)):
                dateList.append(queryResult[i][0])
            return list(set(dateList))
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close()   

def createSalesXDate(date, subsidiary):
    
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            salesXDate = 'EXEC sp_salesXDate ?, ? , ?'
            cursor.execute(salesXDate,(subsidiary, date,0))
            queryResult = cursor.fetchall()
            salesXDateList = queryResult                              
            if salesXDateList :
                saleXDateDoc(date, subsidiary, salesXDateList)
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection.close()  

def getPaymentType():
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            queryGetDates = 'EXEC sp_getPaymentType ?'
            cursor.execute(queryGetDates,(0))
            queryResult = cursor.fetchall()
            paymentTypeList = []
            for i in range(0,len(queryResult)):
                paymentTypeList.append(queryResult[i][0])
            return paymentTypeList
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    
    finally:
        dbConnection.close() 

def createSalesXDateXPaymentType(date, paymentType, subsidiary):
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            salesXDateXPaymentType = 'EXEC sp_salesXDateXPaymentType ? ,? , ? , ?'
            cursor.execute(salesXDateXPaymentType,(subsidiary, date, paymentType,0))
            queryResult = cursor.fetchall()
            salesXDateXPaymentTypeList = queryResult                              
            if salesXDateXPaymentTypeList :
                saleXDatexPaymentTypeDoc(salesXDateXPaymentTypeList)
            
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection.close()  

def createSalesXPaymentType(paymentType, subsidiary):
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            salesXPaymentType = 'EXEC sp_salesXPaymentType ? ,? , ?'
            cursor.execute(salesXPaymentType,(subsidiary ,paymentType,0))
            queryResult = cursor.fetchall()
            salesXPaymentTypeList = queryResult                              
            if salesXPaymentTypeList :
                saleXPaymentTypeDoc(salesXPaymentTypeList)
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection.close()         

def mainBatteryReport():
    
    createSalesXCountry()
    
    document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
    document.add_heading('Ventas por subsidiaria', 0)
    
    subsidiaryList = subsidiaryXCountry()        
    for subsidiary in range(0,len(subsidiaryList)):
        #print(subsidiaryList[subsidiary])
        #print()
        
        createSalesXSubisidiary(subsidiary+1 , subsidiaryList[subsidiary] ) #This function creates the sales per subsidiary report
        productXSubsidiaryList = productsXSubsidiary(subsidiary+1)  
        document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
        document.add_heading('Ventas por subsidiaria por producto', 0)
        for product in range(0,len(productXSubsidiaryList)):
            createSaleXSubsidiaryXProduct(product+1,subsidiary+1 ,productXSubsidiaryList[product])

def saleXSubsidiaryXDateXPaymentType():
    
    document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
    document.add_paragraph("")
    document.add_heading('Ventas por fecha y/o tipo de pago', 0)
    document.save('Reporte de ventas.docx')
    
    subsidiaryList = subsidiaryXCountry()        
    for subsidiary in subsidiaryList:
        tranDateList = getDate(subsidiary)
        newSubsidiary = True
        for date in tranDateList:
            paymentTypeList = getPaymentType()
            
            if newSubsidiary:      
                createSalesXPaymentType(1, subsidiary)
                createSalesXPaymentType(2, subsidiary)
                newSubsidiary = False
            createSalesXDate(date, subsidiary)
            for type in range(0,len(paymentTypeList)):    
                createSalesXDateXPaymentType(date,type+1, subsidiary)

def bestSellerXCountry():

    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
            document.add_paragraph("")
            bestSellerProduct = 'EXEC sp_titleProductTransaction  ?'
            cursor.execute(bestSellerProduct,(0))
            queryResult = cursor.fetchall()
            
            employeeDict = {}
            for item in  range(0,len(queryResult)):
                if queryResult[item][0] not in employeeDict.keys():
                    employeeDict[queryResult[item][0]] = 1
                else:
                    employeeDict[queryResult[item][0]] += 1
            employeeDict = sorted(employeeDict.items(), key=lambda x: x[1], reverse=True)
            
            
            document.add_heading('Los productos más vendidos en ' + country + ' son: ', 1)
            document.add_paragraph("")
            
            document.add_paragraph(employeeDict[0][0])
            document.add_paragraph(employeeDict[1][0])
            document.add_paragraph(employeeDict[2][0])
            
            document.add_heading('Los productos menos vendidos en ' + country + ' son: ', 1)
            document.add_paragraph("")
            
            document.add_paragraph(employeeDict[-1][0])
            document.add_paragraph(employeeDict[-2][0])
            document.add_paragraph(employeeDict[-3][0])
            
            document.save('Reporte de ventas.docx')
                   
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection.close()  

def bestSellerXSubsidiary():
    subsidiaryList = subsidiaryXCountry()        
    for subsidiary in subsidiaryList:
        dbConnection = connectToDatabase(country)
        try:
            with dbConnection.cursor() as cursor:
                document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
                bestSellerProduct = 'EXEC sp_titleProductTransactionXSubsidiary ? ,  ?'
                cursor.execute(bestSellerProduct,(subsidiary,0))
                queryResult = cursor.fetchall()
                
                employeeDict = {}
                for item in  range(0,len(queryResult)):
                    if queryResult[item][0] not in employeeDict.keys():
                        employeeDict[queryResult[item][0]] = 1
                    else:
                        employeeDict[queryResult[item][0]] += 1
                employeeDict = sorted(employeeDict.items(), key=lambda x: x[1], reverse=True)
                
                document.add_heading('El producto más vendidos en ' + subsidiary + ' es: ', 1)
                document.add_paragraph("")
                
                document.add_paragraph(employeeDict[0][0])
                
                document.add_heading('El producto menos vendido en ' + subsidiary + ' es: ', 1)
                document.add_paragraph("")
                
                document.add_paragraph(employeeDict[-1][0])
                
                document.save('Reporte de ventas.docx')
                
                
                    
        except Exception as e:
            print(e)
            return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
        finally:
            dbConnection.close()    

def bestEmployeeXCountry():
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
            document.add_paragraph("")
            document.add_heading('Reportes de Empleado', 0)
            bestSellerProduct = 'EXEC sp_bestEmployee  ?'
            cursor.execute(bestSellerProduct,(0))
            queryResult = cursor.fetchall()
            
            employeeDict = {}
            for item in  range(0,len(queryResult)):
                if queryResult[item][0] not in employeeDict.keys():
                    employeeDict[queryResult[item][0]] = 1
                else:
                    employeeDict[queryResult[item][0]] += 1
            employeeDict = sorted(employeeDict.items(), key=lambda x: x[1], reverse=True)
            
            
            document.add_heading('Los empleados que realizaron más ventas en ' + country + ' son: ', 1)
            document.add_paragraph("")
            
            document.add_paragraph(employeeDict[0][0])
            document.add_paragraph(employeeDict[1][0])
            document.add_paragraph(employeeDict[2][0])
            
            document.add_heading('Los empleados que realizaron menos ventas en ' + country + ' son: ', 1)
            document.add_paragraph("")
            
            document.add_paragraph(employeeDict[-1][0])
            document.add_paragraph(employeeDict[-2][0])
            document.add_paragraph(employeeDict[-3][0])
            
            document.save('Reporte de ventas.docx')
                   
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection.close()  

def bestSellerGlobal():
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
            document.add_paragraph("")
            document.add_heading('Productos más/menos comprados', 0)
            bestSellerProduct = 'EXEC sp_BestSold  ?'
            cursor.execute(bestSellerProduct,(0))
            queryResult = cursor.fetchall()
            
            document.add_heading('El producto más vendido globalmente es: ', 1)
            document.add_paragraph("")
            
            document.add_paragraph(queryResult[0][0])
            
            document.save('Reporte de ventas.docx')
                   
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection.close()  

def worstSellerGlobal():
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
            document.add_paragraph("")
            bestSellerProduct = 'EXEC sp_WorstSold  ?'
            cursor.execute(bestSellerProduct,(0))
            queryResult = cursor.fetchall()
            
            document.add_heading('El producto menos vendido globalmente es: ', 1)
            document.add_paragraph("")
            
            document.add_paragraph(queryResult[0][0])
            
            document.save('Reporte de ventas.docx')
                   
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
    finally:
        dbConnection.close()
        
@app.route('/complain',methods=['GET','POST'])
def complain():
    
    #Get subsidiaries
    dbConnection = connectToDatabase(country)
    try:
        with dbConnection.cursor() as cursor:
            querySubisidiary = 'SELECT S.Title FROM dbo.Subsidiary AS S'
            cursor.execute(querySubisidiary)
            queryResult = cursor.fetchall()
            subsidiaries = []
            for subsidiary in queryResult:
                subsidiaries.append(subsidiary[0])
    except Exception as e:
        print(e)
        return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'

    finally:
        dbConnection.close()
        
    return render_template('quejas.html',**locals())



def bestEmployeeXSubsidiary():
    subsidiaryList = subsidiaryXCountry()        
    for subsidiary in subsidiaryList:
        dbConnection = connectToDatabase(country)
        try:
            with dbConnection.cursor() as cursor:
                document = Document('C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx')
                bestSellerProduct = 'EXEC sp_bestEmployeeXSubsidiary ? ,  ?'
                cursor.execute(bestSellerProduct,(subsidiary,0))
                queryResult = cursor.fetchall()
                
                employeeDict = {}
                for item in  range(0,len(queryResult)):
                    if queryResult[item][0] not in employeeDict.keys():
                        employeeDict[queryResult[item][0]] = 1
                    else:
                        employeeDict[queryResult[item][0]] += 1
                employeeDict = sorted(employeeDict.items(), key=lambda x: x[1], reverse=True)
                document.add_heading('El empleado que realizo más ventas en ' + subsidiary + ' es: ', 1)
                document.add_paragraph("")
                
                document.add_paragraph(employeeDict[0][0])
                
                document.add_heading('El empleado que realizo menos ventas en ' + subsidiary + ' es: ', 1)
                document.add_paragraph("")
                
                document.add_paragraph(employeeDict[-1][0])
                
                document.save('Reporte de ventas.docx')
                
                
                    
        except Exception as e:
            print(e)
            return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
        finally:
            dbConnection.close()   

@app.route('/doComplain',methods=['GET','POST'])
def doComplain():
    subsidiary = request.form['subsidiary']
    employee = request.form['employee']
    suggestion = request.form['suggestion']
    typeOf = request.form['typeOf']
    suggestInformation = {'Subsidiaria':subsidiary,
                          'Empleado':employee,
                          'Usuario': username,
                          'Tipo de recomendación':typeOf,
                          'Comentario':suggestion}
    client = MongoClient('localhost')
    dbCS = client['CustomerService']
    reports = dbCS['Suggestions']
    reports.insert(suggestInformation)
    if employee!="":
        dbConnection = connectToDatabase(country)
        try:
            with dbConnection.cursor() as cursor:
                query = 'SELECT E.Email FROM dbo.Employees AS E WHERE E.Title = ?'
                cursor.execute(query,(employee))
                queryResult = cursor.fetchall()
                email = queryResult[0][0]
                message = ""
                for key in suggestInformation:
                    message+=(str(key) + ":" + str(suggestInformation[key]) + ".\n")
                
                sendEmail(email,message,'Comentario')
                
                
                    
        except Exception as e:
            print(e)
            return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
        finally:
            dbConnection.close()
    else:
        return str('No hubo empleado')
        
        
@app.route('/beginSignUp',methods=['GET','POST'])
def signUp():
        
    return render_template('signUp.html')

@app.route('/ValidateSignUp',methods=['GET','POST'])
def validateSignUp():
    username = request.form['Username']
    keyword = request.form['Keyword']
    fullname = request.form['Name']
    email = request.form['Email']
    phone = request.form['Phone']
    age = request.form['Age']
    country = request.form['countrySelected']
    loc = Nominatim(user_agent="GetLoc") 
    getLocation = loc.geocode(request.form['Address'])
    
    currentDateTime = datetime.datetime.now()
    date = currentDateTime.date()
    year = date.strftime("%Y")
    
    years= age.split("-")
    ageUser = int(year) - int(years[0]) 
    
    if ageUser>= 18:
        address = str(getLocation.latitude) + ' ' +str(getLocation.longitude)
        address = 'POINT(' + address + ')'
        dbConnection = connectToDatabase(country)
        try:
            with dbConnection.cursor() as cursor:
                bestSellerProduct = 'EXEC sp_SignUp ? , ?, ? , ?, ? , ?, ? , ?, ?'
                cursor.execute(bestSellerProduct,(fullname,email,phone,address,ageUser,username,keyword,'Consultant',0))
                queryResult = cursor.fetchall()
                validUser = queryResult[0][0]

                if validUser != 1:
                    return render_template('login.html')
                else:
                    return '''
                    <!DOCTYPE html>
                    <html>
                        <head>
                            <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                            <link href="./static/css/admin.css" rel="stylesheet" />
                            <title>Modificar Inventario</title>
                        </head>
                        <div class="window-notice" id="window-notice" >
                            <div class="content">
                                <div class="content-text">El nombre de usuario ya existe. Ingrese uno nuevo
                                </div>
                                <div class="content-buttons"><a href="/beginSignUp" id="close-button">Aceptar</a></div>
                            </div>
                        </div>
                        <script>
                                    let close_button = document.getElementById('close-button');
                                        close_button.addEventListener("click", function(e) {
                                        e.preventDefault();
                                        document.getElementById("window-notice").style.display = "none";
                                        window.location.href="/beginSignUp";
                                    });
                        </script>
                        '''
                
        except Exception as e:
            return str(e) + 'Exception error. <a href="/">Intente de nuevo.</a>'
        
        finally:
            dbConnection.close()  
        
    else:
         return '''
                    <!DOCTYPE html>
                    <html>
                        <head>
                            <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                            <link href="./static/css/admin.css" rel="stylesheet" />
                            <title>Modificar Inventario</title>
                        </head>
                        <div class="window-notice" id="window-notice" >
                            <div class="content">
                                <div class="content-text">La edad no cumple con la normativa de la tienda.
                                </div>
                                <div class="content-buttons"><a href="/beginSignUp" id="close-button">Aceptar</a></div>
                            </div>
                        </div>
                        <script>
                                    let close_button = document.getElementById('close-button');
                                        close_button.addEventListener("click", function(e) {
                                        e.preventDefault();
                                        document.getElementById("window-notice").style.display = "none";
                                        window.location.href="/beginSignUp";
                                    });
                        </script>
                        '''
        

@app.route('/createReport',methods=['GET','POST'])
def createReport():
    mainBatteryReport()
    saleXSubsidiaryXDateXPaymentType()
    bestSellerGlobal()
    worstSellerGlobal()
    bestSellerXCountry()
    bestSellerXSubsidiary()
    bestEmployeeXCountry()
    bestEmployeeXSubsidiary()
   
    return '''
                <!DOCTYPE html>
                <html>
                    <head>
                        <link href="https://fonts.googleapis.com/css?family=Inter&display=swap" rel="stylesheet" />
                        <link href="./static/css/admin.css" rel="stylesheet" />
                        <title>Modificar Inventario</title>
                    </head>
                    <div class="window-notice" id="window-notice" >
                        <div class="content">
                            <div class="content-text">Documento de reportes creado con exito.
                            </div>
                            <div class="content-buttons"><a href="#" id="close-button">Aceptar</a></div>
                        </div>
                    </div>
                    <script>
                                let close_button = document.getElementById('close-button');
                                    close_button.addEventListener("click", function(e) {
                                    e.preventDefault();
                                    document.getElementById("window-notice").style.display = "none";
                                    window.location.href="/returnAdmin";
                                });
                    </script>
                    '''
    
#'C:/Users/luist/OneDrive/Escritorio/LiqueursStore/Reporte de ventas.docx'
#'C:/Users/Sebastian/Desktop/TEC/IVSemestre/BasesDatosII/Proyecto/LiqueursStore/Reporte de ventas.docx'

#Run application
if __name__ == '__main__':
    app.run(host="0.0.0.0",port=4000,debug=True)